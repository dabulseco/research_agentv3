import os
import warnings

# CRITICAL: Set environment variables BEFORE importing crewai
os.environ['CREWAI_TELEMETRY_ENABLED'] = 'false'
os.environ['OTEL_SDK_DISABLED'] = 'true'
# Set dummy OpenAI key to prevent CrewAI from complaining
os.environ['OPENAI_API_KEY'] = 'sk-dummy-key-for-local-llm'
os.environ['OPENAI_MODEL_NAME'] = 'gpt-4'  # Dummy model name
# Disable interactive prompts
os.environ['CREWAI_DISABLE_TELEMETRY_PROMPT'] = 'true'
os.environ['CREWAI_STORAGE_DIR'] = './.crewai_storage'

# Suppress deprecation warnings
warnings.filterwarnings('ignore', category=DeprecationWarning)

import streamlit as st
from crewai import Agent, Task, Crew, Process

try:
    from langchain_ollama import OllamaLLM
except ImportError:
    # Fallback to old import if langchain-ollama not installed
    from langchain_community.llms import Ollama as OllamaLLM

from crewai.tools import BaseTool
from langchain_core.tools import Tool
from duckduckgo_search import DDGS
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options as ChromeOptions
from selenium.webdriver.firefox.options import Options as FirefoxOptions
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
from webdriver_manager.firefox import GeckoDriverManager
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.firefox.service import Service as FirefoxService
import json
from datetime import datetime
import re
import io
from bs4 import BeautifulSoup
import requests
import subprocess
import time
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Initialize session state
if 'workflow_stage' not in st.session_state:
    st.session_state.workflow_stage = 0
if 'research_plan' not in st.session_state:
    st.session_state.research_plan = None
if 'initial_research' not in st.session_state:
    st.session_state.initial_research = None
if 'gap_analysis' not in st.session_state:
    st.session_state.gap_analysis = None
if 'web_research' not in st.session_state:
    st.session_state.web_research = None
if 'blog_post' not in st.session_state:
    st.session_state.blog_post = None
if 'user_input' not in st.session_state:
    st.session_state.user_input = ""
if 'browser_type' not in st.session_state:
    st.session_state.browser_type = "chrome"
if 'search_method' not in st.session_state:
    st.session_state.search_method = "selenium"

# Research context configuration
if 'perspective' not in st.session_state:
    st.session_state.perspective = ""
if 'audience' not in st.session_state:
    st.session_state.audience = ""
if 'framing_notes' not in st.session_state:
    st.session_state.framing_notes = ""
if 'context_configured' not in st.session_state:
    st.session_state.context_configured = False

# Output format and length
if 'output_format' not in st.session_state:
    st.session_state.output_format = "blog_post"
if 'output_length' not in st.session_state:
    st.session_state.output_length = "~1,000 words (2-3 pages)"
if 'output_format_selected' not in st.session_state:
    st.session_state.output_format_selected = False
if 'written_output' not in st.session_state:
    st.session_state.written_output = None

# Custom resources (URLs and PDFs)
if 'custom_resources' not in st.session_state:
    st.session_state.custom_resources = []  # list[dict] with keys: type, source, title, content
if 'resource_added_after_stage' not in st.session_state:
    st.session_state.resource_added_after_stage = None

# ---------------------------------------------------------------------------
# Output format templates
# ---------------------------------------------------------------------------
FORMAT_TEMPLATES: dict[str, dict[str, str]] = {
    "blog_post": {
        "label": "Blog Post",
        "agent_role": "Science Communicator and Technical Writer",
        "agent_goal": "Create an engaging, well-structured blog post with proper citations",
        "agent_backstory": (
            "You are an expert science communicator who makes complex topics "
            "accessible and compelling for a broad readership."
        ),
        "structure_guidance": (
            "Write with an engaging title, a hook introduction, logically ordered "
            "sections with clear headers, and a conclusion that summarises key points. "
            "Use a conversational but authoritative tone. Include source citations."
        ),
    },
    "research_paper": {
        "label": "Research Paper",
        "agent_role": "Academic Research Writer",
        "agent_goal": "Produce a rigorous, well-cited academic-style research paper",
        "agent_backstory": (
            "You are an experienced academic writer who synthesises research into "
            "formal papers suitable for peer-reviewed publication."
        ),
        "structure_guidance": (
            "Follow academic structure: Abstract, Introduction, Background / Literature "
            "Review, Methods / Approach, Findings / Results, Discussion, Conclusion, "
            "References. Use formal tone, third person, and precise language."
        ),
    },
    "essay": {
        "label": "Essay",
        "agent_role": "Essayist and Analytical Writer",
        "agent_goal": "Write a thoughtful, well-argued analytical essay",
        "agent_backstory": (
            "You are a skilled essayist who builds persuasive arguments supported "
            "by evidence and careful reasoning."
        ),
        "structure_guidance": (
            "Open with a clear thesis statement, develop the argument through "
            "evidence-backed body paragraphs, address counter-arguments, and close "
            "with a strong conclusion that reinforces the thesis."
        ),
    },
    "executive_summary": {
        "label": "Executive Summary / Briefing",
        "agent_role": "Policy Analyst and Strategic Communications Expert",
        "agent_goal": "Produce a concise, decision-oriented executive summary",
        "agent_backstory": (
            "You brief senior stakeholders, executives, and policymakers with "
            "clear, actionable summaries that enable rapid informed decisions."
        ),
        "structure_guidance": (
            "Use sections: Key Findings, Background, Analysis, Recommendations, "
            "Conclusion. Lead with the most important insights. Use bullet points "
            "and concise headers. Avoid jargon; favour clarity."
        ),
    },
    "nsf_proposal": {
        "label": "NSF-Style Grant Proposal",
        "agent_role": "Grant Writing Specialist",
        "agent_goal": "Write a compelling NSF-style research proposal that demonstrates intellectual merit and broader impacts",
        "agent_backstory": (
            "You are an experienced grant writer who has successfully authored "
            "multiple NSF proposals. You know how to frame research to emphasise "
            "novelty, significance, and societal benefit."
        ),
        "structure_guidance": (
            "Use flexible NSF-like sections appropriate to the research topic — "
            "for example: Project Summary, Introduction / Motivation, "
            "Background and Related Work, Research Objectives, Proposed Approach, "
            "Intellectual Merit, Broader Impacts, Timeline, References. "
            "Emphasise innovation, feasibility, and clear outcomes."
        ),
    },
    "explainer": {
        "label": "Explainer / Educational Article",
        "agent_role": "Science Educator and Explainer Writer",
        "agent_goal": "Create a clear, accessible educational explainer for learners",
        "agent_backstory": (
            "You excel at breaking down complex topics for learners, using plain "
            "language, relatable analogies, and well-chosen examples."
        ),
        "structure_guidance": (
            "Start with the big picture and why it matters, then build complexity "
            "gradually. Use examples and analogies. End with key takeaways or a "
            "'What to remember' box. Avoid unexplained jargon."
        ),
    },
}

# Section lists used by step5_write_output's section-by-section mode.
# Each entry maps an output_format key to an ordered list of section names.
# These names appear as spinner labels and as ## headings in the output.
TEMPLATE_SECTIONS: dict[str, list[str]] = {
    "blog_post": [
        "Title and Introduction",
        "Background and Context",
        "Main Analysis — Part 1",
        "Main Analysis — Part 2",
        "Key Insights and Implications",
        "Conclusion and Summary",
    ],
    "research_paper": [
        "Abstract",
        "Introduction",
        "Background and Literature Review",
        "Methods and Approach",
        "Findings and Results",
        "Discussion",
        "Conclusion",
        "References",
    ],
    "essay": [
        "Introduction and Thesis",
        "Background and Context",
        "Main Argument — Part 1",
        "Main Argument — Part 2",
        "Counter-arguments and Rebuttals",
        "Conclusion",
    ],
    "executive_summary": [
        "Executive Overview",
        "Key Findings",
        "Background and Analysis",
        "Recommendations",
        "Conclusion",
    ],
    "nsf_proposal": [
        "Project Summary",
        "Introduction and Motivation",
        "Background and Related Work",
        "Research Objectives",
        "Proposed Approach",
        "Intellectual Merit and Broader Impacts",
        "Timeline and References",
    ],
    "explainer": [
        "Introduction — The Big Picture",
        "Core Concepts Explained",
        "Deep Dive — Part 1",
        "Deep Dive — Part 2",
        "Real-world Applications and Examples",
        "Key Takeaways",
    ],
}

# ---------------------------------------------------------------------------
# Context and resource helpers
# ---------------------------------------------------------------------------

def build_context_block() -> str:
    """Return a formatted context string to inject into all agent prompts."""
    parts: list[str] = []
    if st.session_state.perspective:
        parts.append(f"PERSPECTIVE/VOICE: {st.session_state.perspective}")
    if st.session_state.audience:
        parts.append(f"TARGET AUDIENCE: {st.session_state.audience}")
    if st.session_state.framing_notes:
        parts.append(f"FRAMING NOTES: {st.session_state.framing_notes}")
    if not parts:
        return ""
    return "\n\nCONTEXT CONFIGURATION (apply throughout all research and writing):\n" + "\n".join(parts) + "\n"


def build_resources_block() -> str:
    """Return formatted custom resources for injection into agent prompts."""
    resources: list[dict] = st.session_state.get('custom_resources', [])
    if not resources:
        return ""
    lines: list[str] = [
        "\n\nCUSTOM RESOURCES PROVIDED BY USER "
        "(treat as authoritative primary sources; cite them explicitly in output):"
    ]
    for i, r in enumerate(resources, 1):
        lines.append(
            f"\n--- Resource {i}: {r['title']} "
            f"({r['type'].upper()}: {r['source']}) ---"
        )
        # Per-resource truncation is a safety net; main truncation happens at ingest
        lines.append(r['content'][:3000])
        lines.append("--- End Resource ---")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# URL scraping helper
# ---------------------------------------------------------------------------

def scrape_url(url: str) -> dict[str, str]:
    """Scrape a URL and return a {title, content} dict."""
    try:
        headers = {"User-Agent": "Mozilla/5.0 (compatible; ResearchAgent/1.0)"}
        resp = requests.get(url, headers=headers, timeout=15)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        title: str = (
            soup.title.string.strip()
            if soup.title and soup.title.string
            else url
        )
        text: str = soup.get_text(separator="\n", strip=True)
        words = text.split()
        if len(words) > 4000:
            text = " ".join(words[:4000]) + "\n[... content truncated ...]"
        return {"title": title, "content": text}
    except Exception as exc:
        return {"title": url, "content": f"[Could not scrape URL: {exc}]"}


# ---------------------------------------------------------------------------
# PDF extraction helper (chunking for large documents)
# ---------------------------------------------------------------------------

def extract_pdf(file_bytes: bytes, filename: str) -> dict[str, str]:
    """Extract text from a PDF; chunk large documents into sections."""
    try:
        import pdfplumber
        import io as _io

        text_parts: list[str] = []
        with pdfplumber.open(_io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        full_text = "\n".join(text_parts)
        words = full_text.split()

        if len(words) > 4000:
            # Split into ~1 000-word chunks (up to 20 chunks = ~20 000 words)
            chunks: list[str] = []
            for start in range(0, min(len(words), 20_000), 1_000):
                chunks.append(" ".join(words[start : start + 1_000]))
            content = "\n\n[SECTION BREAK]\n\n".join(chunks)
            content += "\n[Large document — first ~20,000 words shown in sections]"
        else:
            content = full_text

        return {"title": filename, "content": content}
    except Exception as exc:
        return {"title": filename, "content": f"[Could not extract PDF: {exc}]"}


# ---------------------------------------------------------------------------
# Resource panel sidebar renderer (called inside the sidebar context)
# ---------------------------------------------------------------------------

def render_resource_panel() -> None:
    """Render the persistent Research Resources panel in the sidebar."""
    st.markdown("---")
    st.markdown("### 📎 Research Resources")
    st.caption("Add URLs or PDFs at any stage. They will be injected into all remaining research and writing steps.")

    # --- URL input ---
    url_input = st.text_input(
        "Add a URL",
        key="url_input_field",
        placeholder="https://example.com/article",
    )
    if st.button("➕ Add URL", key="add_url_btn"):
        raw = url_input.strip()
        if raw:
            already = [r["source"] for r in st.session_state.custom_resources]
            if raw in already:
                st.warning("This URL is already in the resource list.")
            else:
                with st.spinner(f"Scraping {raw} …"):
                    result = scrape_url(raw)
                st.session_state.custom_resources.append({
                    "type": "url",
                    "source": raw,
                    "title": result["title"],
                    "content": result["content"],
                })
                # Record that a resource was added at this workflow stage
                stage = st.session_state.workflow_stage
                if stage > 0:
                    st.session_state.resource_added_after_stage = stage
                st.success(f"Added: {result['title']}")
                st.rerun()
        else:
            st.warning("Please enter a URL.")

    # --- PDF upload ---
    pdf_file = st.file_uploader(
        "Upload a PDF",
        type=["pdf"],
        key="pdf_uploader",
        help="Large PDFs are chunked automatically.",
    )
    if pdf_file is not None:
        already_titles = [r["source"] for r in st.session_state.custom_resources]
        if pdf_file.name not in already_titles:
            if st.button("➕ Add PDF", key="add_pdf_btn"):
                with st.spinner(f"Extracting text from {pdf_file.name} …"):
                    result = extract_pdf(pdf_file.read(), pdf_file.name)
                st.session_state.custom_resources.append({
                    "type": "pdf",
                    "source": pdf_file.name,
                    "title": result["title"],
                    "content": result["content"],
                })
                stage = st.session_state.workflow_stage
                if stage > 0:
                    st.session_state.resource_added_after_stage = stage
                st.success(f"Added PDF: {pdf_file.name}")
                st.rerun()

    # --- Resource list ---
    if st.session_state.custom_resources:
        st.markdown("**Current Resources:**")
        for idx, r in enumerate(st.session_state.custom_resources):
            col_label, col_btn = st.columns([4, 1])
            icon = "🔗" if r["type"] == "url" else "📄"
            col_label.markdown(f"{icon} {r['title'][:40]}")
            if col_btn.button("✕", key=f"remove_resource_{idx}"):
                st.session_state.custom_resources.pop(idx)
                st.rerun()

    # --- Re-run buttons when resource added after a completed stage ---
    stage_now = st.session_state.workflow_stage
    added_at = st.session_state.resource_added_after_stage

    if added_at is not None and stage_now > 1:
        st.markdown("---")
        st.warning("A resource was added after research had already progressed. Re-run from a prior stage to incorporate it.")

        # Offer the most useful re-run points based on current stage
        if stage_now >= 5:  # written output exists
            if st.button("🔁 Re-run from Output Writing", key="rerun_writing"):
                st.session_state.written_output = None
                st.session_state.blog_post = None

                st.session_state.workflow_stage = 45
                st.session_state.resource_added_after_stage = None
                st.rerun()
        if stage_now >= 4 or stage_now == 45:  # web research done
            if st.button("🔁 Re-run from Web Research", key="rerun_web"):
                st.session_state.web_research = None
                st.session_state.written_output = None
                st.session_state.blog_post = None

                st.session_state.workflow_stage = 3
                st.session_state.resource_added_after_stage = None
                st.rerun()
        if stage_now >= 3:  # gap analysis done
            if st.button("🔁 Re-run from Gap Analysis", key="rerun_gap"):
                st.session_state.gap_analysis = None
                st.session_state.web_research = None
                st.session_state.written_output = None
                st.session_state.blog_post = None

                st.session_state.workflow_stage = 2
                st.session_state.resource_added_after_stage = None
                st.rerun()
        if stage_now >= 2:  # initial research done
            if st.button("🔁 Re-run from Initial Research", key="rerun_initial"):
                st.session_state.research_plan = None
                st.session_state.initial_research = None
                st.session_state.gap_analysis = None
                st.session_state.web_research = None
                st.session_state.written_output = None
                st.session_state.blog_post = None

                st.session_state.workflow_stage = 1
                st.session_state.resource_added_after_stage = None
                st.rerun()


# Configure page
st.set_page_config(page_title="AI Research Agent", layout="wide")

# Function to get locally installed Ollama models
def get_installed_ollama_models():
    """Fetch list of locally installed Ollama models"""
    try:
        result = subprocess.run(
            ['ollama', 'list'],
            capture_output=True,
            text=True,
            timeout=10
        )
        
        if result.returncode == 0:
            lines = result.stdout.strip().split('\n')
            # Skip the header line and extract model names
            models = []
            for line in lines[1:]:  # Skip header
                if line.strip():
                    # Model name is the first column (includes tag like gemma3:27b)
                    parts = line.split()
                    if parts:
                        model_name = parts[0]  # Keep full name with tag
                        if model_name and model_name not in models:
                            models.append(model_name)
            
            if not models:
                st.error("No Ollama models found! Please install a model first.")
                st.code("ollama pull llama3.2", language="bash")
                return []
            return models
        else:
            st.warning("Could not fetch Ollama models. Make sure Ollama is running.")
            return []
    except FileNotFoundError:
        st.error("Ollama not found. Please install Ollama from https://ollama.ai")
        return []
    except subprocess.TimeoutExpired:
        st.warning("Ollama command timed out. Using default models.")
        return []
    except Exception as e:
        st.warning(f"Error fetching models: {str(e)}")
        return []

# Verify model exists in Ollama
def verify_model_exists(model_name):
    """Check if a model exists in Ollama"""
    try:
        result = subprocess.run(
            ['ollama', 'list'],
            capture_output=True,
            text=True,
            timeout=10
        )
        
        if result.returncode == 0:
            # Check if model_name appears exactly in the output
            lines = result.stdout.strip().split('\n')
            for line in lines[1:]:  # Skip header
                if line.strip():
                    parts = line.split()
                    if parts and parts[0] == model_name:
                        return True
            return False
        return False
    except Exception:
        return False

# Initialize Ollama LLM
@st.cache_resource
def get_llm(model_name: str = "llama2") -> OllamaLLM:
    """Standard LLM for research/analysis agents (moderate context window)."""
    return OllamaLLM(
        model=model_name,
        base_url="http://localhost:11434",
        temperature=0.7,
        num_ctx=8192,       # Up from Ollama's implicit 2048 default
        num_predict=2048,   # Allow non-writer agents to produce longer responses
    )


# Selenium-based web search
class SeleniumSearcher:
    """Headless browser-based web search using Selenium"""
    
    def __init__(self, browser_type="chrome"):
        self.browser_type = browser_type
        self.driver = None
        self.is_mac_arm = self._check_mac_arm()
    
    def _check_mac_arm(self):
        """Check if running on Mac ARM (M1/M2/M3)"""
        import platform
        return platform.system() == 'Darwin' and platform.machine() == 'arm64'
    
    def _init_driver(self):
        """Initialize the webdriver"""
        try:
            if self.browser_type == "chrome":
                chrome_options = ChromeOptions()
                chrome_options.add_argument('--headless=new')  # New headless mode
                chrome_options.add_argument('--no-sandbox')
                chrome_options.add_argument('--disable-dev-shm-usage')
                chrome_options.add_argument('--disable-gpu')
                chrome_options.add_argument('--window-size=1920,1080')
                chrome_options.add_argument('--user-agent=Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36')
                chrome_options.add_argument('--disable-blink-features=AutomationControlled')
                chrome_options.add_experimental_option("excludeSwitches", ["enable-automation"])
                chrome_options.add_experimental_option('useAutomationExtension', False)
                
                # MacOS ARM-specific settings
                if self.is_mac_arm:
                    chrome_options.add_argument('--disable-software-rasterizer')
                    chrome_options.add_argument('--disable-extensions')
                    # Don't use sandbox on Mac ARM due to compatibility issues
                    chrome_options.add_argument('--no-sandbox')
                
                try:
                    service = ChromeService(ChromeDriverManager().install())
                    self.driver = webdriver.Chrome(service=service, options=chrome_options)
                except Exception as e:
                    # Fallback: try without service
                    st.warning(f"ChromeDriver manager failed, trying system Chrome: {str(e)}")
                    self.driver = webdriver.Chrome(options=chrome_options)
                    
            else:  # firefox
                firefox_options = FirefoxOptions()
                firefox_options.add_argument('--headless')
                firefox_options.add_argument('--width=1920')
                firefox_options.add_argument('--height=1080')
                
                # MacOS-specific Firefox settings
                if self.is_mac_arm:
                    firefox_options.set_preference('media.navigator.enabled', False)
                    firefox_options.set_preference('media.peerconnection.enabled', False)
                
                try:
                    service = FirefoxService(GeckoDriverManager().install())
                    self.driver = webdriver.Firefox(service=service, options=firefox_options)
                except Exception as e:
                    # Fallback: try without service
                    st.warning(f"GeckoDriver manager failed, trying system Firefox: {str(e)}")
                    self.driver = webdriver.Firefox(options=firefox_options)
            
            # Set page load timeout
            self.driver.set_page_load_timeout(30)
            return True
            
        except Exception as e:
            st.error(f"Failed to initialize {self.browser_type} driver: {str(e)}")
            if self.is_mac_arm:
                st.info("💡 **MacOS M1/M2 Tip**: If Chrome fails, try Firefox or install Chrome for ARM from https://www.google.com/chrome/")
            return False
    
    def search_google(self, query, max_results=5):
        """Search using Google"""
        if not self.driver and not self._init_driver():
            return []
        
        results = []
        try:
            # Navigate to Google
            search_url = f"https://www.google.com/search?q={requests.utils.quote(query)}"
            self.driver.get(search_url)
            
            # Wait for results to load (with timeout)
            try:
                WebDriverWait(self.driver, 10).until(
                    EC.presence_of_element_located((By.CSS_SELECTOR, "div.g"))
                )
            except Exception:
                # If wait fails, still try to parse what loaded
                pass
            
            time.sleep(1)  # Small additional wait
            
            # Parse results
            soup = BeautifulSoup(self.driver.page_source, 'html.parser')
            
            # Find search result divs
            search_results = soup.find_all('div', class_='g')
            
            for result in search_results[:max_results]:
                try:
                    # Extract title
                    title_elem = result.find('h3')
                    title = title_elem.get_text() if title_elem else 'No title'
                    
                    # Extract link
                    link_elem = result.find('a')
                    link = link_elem.get('href') if link_elem else ''
                    
                    # Extract snippet
                    snippet_elem = result.find('div', class_=['VwiC3b', 'yXK7lf'])
                    snippet = snippet_elem.get_text() if snippet_elem else 'No snippet available'
                    
                    if title and link:
                        results.append({
                            'title': title,
                            'link': link,
                            'snippet': snippet
                        })
                except Exception:
                    continue
            
        except Exception as e:
            st.warning(f"Google search error: {str(e)}")
        
        return results
    
    def search_bing(self, query, max_results=5):
        """Search using Bing"""
        if not self.driver and not self._init_driver():
            return []
        
        results = []
        try:
            # Navigate to Bing
            search_url = f"https://www.bing.com/search?q={requests.utils.quote(query)}"
            self.driver.get(search_url)
            
            # Wait for results
            try:
                WebDriverWait(self.driver, 10).until(
                    EC.presence_of_element_located((By.CSS_SELECTOR, "li.b_algo"))
                )
            except Exception:
                pass
            
            time.sleep(1)
            
            # Parse results
            soup = BeautifulSoup(self.driver.page_source, 'html.parser')
            
            # Find search result items
            search_results = soup.find_all('li', class_='b_algo')
            
            for result in search_results[:max_results]:
                try:
                    title_elem = result.find('h2')
                    title = title_elem.get_text() if title_elem else 'No title'
                    
                    link_elem = result.find('a')
                    link = link_elem.get('href') if link_elem else ''
                    
                    snippet_elem = result.find('p')
                    snippet = snippet_elem.get_text() if snippet_elem else 'No snippet available'
                    
                    if title and link:
                        results.append({
                            'title': title,
                            'link': link,
                            'snippet': snippet
                        })
                except Exception:
                    continue
            
        except Exception as e:
            st.warning(f"Bing search error: {str(e)}")
        
        return results
    
    def close(self):
        """Close the webdriver"""
        if self.driver:
            try:
                self.driver.quit()
            except Exception:
                pass
            self.driver = None

# Web search tool - CrewAI compatible with multiple backends
class WebSearchTool(BaseTool):
    name: str = "web_search"
    description: str = "Search the web for information. Input should be a search query string. Returns top search results with titles, links, and snippets."
    
    def _run(self, query: str) -> str:
        """Search the web using selected method"""
        search_method = st.session_state.get('search_method', 'selenium')
        browser_type = st.session_state.get('browser_type', 'chrome')
        
        if search_method == 'selenium':
            return self._selenium_search(query, browser_type)
        else:
            return self._ddg_search(query)
    
    def _selenium_search(self, query: str, browser_type: str) -> str:
        """Search using Selenium headless browser"""
        searcher = SeleniumSearcher(browser_type)
        
        try:
            # Try Google first
            results = searcher.search_google(query, max_results=5)
            
            # If Google fails, try Bing
            if not results:
                results = searcher.search_bing(query, max_results=5)
            
            if results:
                return json.dumps(results, indent=2)
            
            return json.dumps([{
                'title': 'No results found',
                'link': '',
                'snippet': f'No web results found for "{query}". Please try a different search query.'
            }], indent=2)
            
        except Exception as e:
            return json.dumps([{
                'title': 'Search error',
                'link': '',
                'snippet': f'Error searching for "{query}": {str(e)}'
            }], indent=2)
        finally:
            searcher.close()
    
    def _ddg_search(self, query: str) -> str:
        """Search using DuckDuckGo with retry logic"""
        for attempt in range(3):
            try:
                results = []
                with DDGS() as ddgs:
                    search_results = list(ddgs.text(query, max_results=5))
                    for r in search_results:
                        results.append({
                            'title': r.get('title', 'No title'),
                            'link': r.get('href', ''),
                            'snippet': r.get('body', 'No snippet available')
                        })
                
                if results:
                    return json.dumps(results, indent=2)
                
                if attempt < 2:
                    time.sleep(2)
                    
            except Exception as e:
                if attempt < 2:
                    time.sleep(2)
                    continue
                return json.dumps([{
                    'title': f'Search temporarily unavailable',
                    'link': '',
                    'snippet': f'Unable to perform web search for "{query}". Error: {str(e)}'
                }], indent=2)
        
        return json.dumps([{
            'title': 'No results found',
            'link': '',
            'snippet': f'No web results found for "{query}".'
        }], indent=2)

# Create search tool instance
search_tool = WebSearchTool()

# Step 1: Interpret Request and Create Research Plan
def step1_interpret_and_plan(user_request: str, llm, context: str = "", resources: str = ""):
    """Interpret user request and create research plan."""
    st.subheader("Step 1: Interpreting Request & Creating Research Plan")

    llm_string = f"ollama/{llm.model}"

    interpreter_agent = Agent(
        role="Research Request Interpreter",
        goal=(
            "Understand the user request and create a comprehensive research plan "
            "that is tailored to the specified audience and perspective."
        ),
        backstory=(
            "You are an expert at understanding research needs and creating "
            "structured research plans. You always tailor your plans to the "
            "intended audience and the researcher's perspective."
        ),
        llm=llm_string,
        verbose=True,
        allow_delegation=False,
    )

    interpret_task = Task(
        description=f"""
        Analyze this research request: "{user_request}"
        {context}
        {resources}

        Create a detailed research plan that includes:
        1. Main research questions (3-5 questions) — framed for the specified audience and perspective
        2. Key topics to investigate
        3. Suggested search queries
        4. Expected information structure

        If custom resources have been provided above, note the topics they cover so later
        steps do not unnecessarily duplicate that research.

        Format your response as a structured JSON with these keys:
        - main_questions: list of questions
        - topics: list of topics
        - search_queries: list of queries
        - structure: outline of expected information
        - covered_by_resources: list of topics already addressed by provided resources (empty list if none)
        """,
        agent=interpreter_agent,
        expected_output="A structured research plan in JSON format",
    )

    crew = Crew(
        agents=[interpreter_agent],
        tasks=[interpret_task],
        process=Process.sequential,
        verbose=True,
        memory=False,
        cache=False,
        output_log_file=False,
    )

    with st.spinner("Creating research plan..."):
        result = crew.kickoff()

    return str(result)

# Step 2: Execute Initial Research with Local LLM
def step2_initial_research(research_plan: str, llm, context: str = "", resources: str = ""):
    """Execute initial research using local LLM knowledge."""
    st.subheader("Step 2: Executing Initial Research (Local LLM)")

    llm_string = f"ollama/{llm.model}"

    researcher_agent = Agent(
        role="Knowledge Researcher",
        goal=(
            "Provide comprehensive information based on existing knowledge, "
            "always keeping the specified audience and perspective in mind."
        ),
        backstory=(
            "You are a knowledgeable researcher who synthesises information on "
            "diverse topics. You adapt the depth and framing of your research to "
            "the intended audience and the researcher's stated perspective."
        ),
        llm=llm_string,
        verbose=True,
        allow_delegation=False,
    )

    research_task = Task(
        description=f"""
        Based on this research plan:
        {research_plan}
        {context}
        {resources}

        Provide detailed information on each topic using your existing knowledge.
        Structure your response with:
        - Clear sections for each main question
        - Detailed explanations pitched at the specified audience
        - Key facts and concepts framed from the specified perspective
        - Note any areas where information might be limited or outdated

        If custom resources are provided above, integrate their content and do not
        repeat research that those resources already cover — instead reference them.
        """,
        agent=researcher_agent,
        expected_output="Comprehensive research report based on existing knowledge",
    )

    crew = Crew(
        agents=[researcher_agent],
        tasks=[research_task],
        process=Process.sequential,
        verbose=True,
        memory=False,
        cache=False,
        output_log_file=False,
    )

    with st.spinner("Conducting initial research..."):
        result = crew.kickoff()

    return str(result)

# Step 3: Gap Analysis
def step3_gap_analysis(
    research_plan: str,
    initial_research: str,
    llm,
    context: str = "",
    resources: str = "",
):
    """Analyze gaps in the initial research."""
    st.subheader("Step 3: Conducting Gap Analysis")

    llm_string = f"ollama/{llm.model}"

    analyst_agent = Agent(
        role="Research Gap Analyst",
        goal=(
            "Identify missing information and areas needing web research, "
            "taking into account the target audience and any resources already provided."
        ),
        backstory=(
            "You excel at identifying gaps in research and determining exactly "
            "what additional information is needed to produce a complete, "
            "audience-appropriate output."
        ),
        llm=llm_string,
        verbose=True,
        allow_delegation=False,
    )

    gap_task = Task(
        description=f"""
        Research Plan:
        {research_plan}

        Initial Research:
        {initial_research}
        {context}
        {resources}

        Analyze the initial research and identify:
        1. Missing information or unanswered questions (considering the audience and perspective above)
        2. Areas where current information might be outdated
        3. Topics requiring recent data or statistics
        4. Specific search queries needed to fill gaps

        IMPORTANT: If custom resources are listed above, do NOT flag the topics they
        cover as gaps — they are already addressed. Only flag genuine missing areas.

        Format as JSON with:
        - gaps: list of identified gaps
        - priority: high/medium/low for each gap
        - search_queries: specific queries to fill each gap
        """,
        agent=analyst_agent,
        expected_output="Gap analysis with prioritized search queries",
    )

    crew = Crew(
        agents=[analyst_agent],
        tasks=[gap_task],
        process=Process.sequential,
        verbose=True,
        memory=False,
        cache=False,
        output_log_file=False,
    )

    with st.spinner("Analyzing research gaps..."):
        result = crew.kickoff()

    return str(result)

# Step 4: Web Research to Fill Gaps
def step4_web_research(gap_analysis: str, llm, context: str = "", resources: str = ""):
    """Conduct web searches to fill identified gaps."""
    st.subheader("Step 4: Conducting Web Research")

    llm_string = f"ollama/{llm.model}"

    web_researcher_agent = Agent(
        role="Web Research Specialist",
        goal=(
            "Find and synthesise information from the web to fill research gaps, "
            "complementing any custom resources already provided. "
            "If web search is unavailable, synthesise available information and "
            "note what could not be verified."
        ),
        backstory=(
            "You are skilled at searching the web and extracting relevant information "
            "for a specific audience and perspective. When searches fail, you document "
            "what information is missing and work with available data."
        ),
        llm=llm_string,
        tools=[search_tool],
        verbose=True,
        allow_delegation=False,
        max_iter=10,
    )

    web_research_task = Task(
        description=f"""
        Gap Analysis:
        {gap_analysis}
        {context}
        {resources}

        Use the web_search tool to find information for the most important identified gaps.
        For each gap:
        1. Try searching with clear, specific queries
        2. If a search fails or returns no results, note this and move to the next gap
        3. Summarise findings from successful searches
        4. Note sources when available

        IMPORTANT:
        - If custom resources are listed above, treat them as already-gathered data —
          do NOT re-search topics they cover; simply reference them.
        - Focus on 2-3 most critical gaps to avoid rate limiting
        - If web search consistently fails, compile findings from searches that did work
        - Document which gaps could not be filled due to search limitations
        - Provide a summary even if some searches fail
        - Frame all findings for the specified audience and perspective

        Compile findings into a report that includes:
        - Information found from successful searches
        - Sources/links when available
        - List of gaps that could not be filled and why
        """,
        agent=web_researcher_agent,
        expected_output="Web research findings with sources, including documentation of any search failures",
    )
    
    crew = Crew(
        agents=[web_researcher_agent],
        tasks=[web_research_task],
        process=Process.sequential,
        verbose=True,
        memory=False,
        cache=False,
        output_log_file=False
    )
    
    with st.spinner("Searching the web for additional information..."):
        try:
            result = crew.kickoff()
            return str(result)
        except Exception as e:
            error_msg = f"Web research encountered an error: {str(e)}\n\nProceeding with available information from earlier research phases."
            st.warning(error_msg)
            return error_msg

def _parse_target_word_count(output_length: str) -> int:
    """Extract a numeric word-count target from the output_length string.

    Examples:
        '~5,000 words (10 pages)' -> 5000
        '~500 words (1 page)'     -> 500
        '3,500 words or 7 pages'  -> 3500
        'Custom...'               -> 1000 (safe fallback)
    """
    cleaned = output_length.replace(",", "")
    m = re.search(r"(\d+)", cleaned)
    return int(m.group(1)) if m else 1000


def _prime_ollama_context(model: str, num_ctx: int, base_url: str = "http://localhost:11434") -> None:
    """Prime Ollama to load the model with the desired num_ctx.

    Sends a minimal /api/generate request with num_ctx in the options block.
    Ollama loads (or re-loads) the model with the specified context window,
    which then persists for subsequent LiteLLM/CrewAI calls until the model
    is unloaded from memory.

    This is the only reliable way to set num_ctx when going through the
    LiteLLM/CrewAI stack, because CrewAI hardcodes litellm.drop_params=True
    which silently discards extra kwargs like num_ctx before the Ollama call.
    """
    try:
        resp = requests.post(
            f"{base_url}/api/generate",
            json={
                "model": model,
                "prompt": ".",
                "stream": False,
                "options": {
                    "num_ctx": num_ctx,
                    "num_predict": 1,   # Generate only 1 token — just to trigger model load
                },
            },
            timeout=60,
        )
        resp.raise_for_status()
    except Exception as e:
        # Non-fatal: log but don't crash. CrewAI will still call the model;
        # it just may use a smaller context window than requested.
        st.warning(f"\u26a0\ufe0f Could not prime Ollama context window ({num_ctx} tokens): {e}")


def _markdown_to_docx(markdown_text: str, title: str = "Research Output") -> bytes:
    """Convert a markdown string to a styled DOCX file and return raw bytes.

    Handles the following markdown constructs:
    - # / ## / ### headings  → Word Heading 1 / 2 / 3 styles
    - **bold** text          → bold run
    - *italic* text          → italic run
    - `code`                 → Courier New monospace run
    - --- (horizontal rule)  → paragraph with bottom border
    - Blank lines            → paragraph breaks
    - Unordered list items (- or *) → List Bullet style
    - Ordered list items (1. 2. …)  → List Number style
    - Everything else        → Normal paragraph with inline formatting
    """
    doc = Document()

    # Page margins — 1 inch all sides
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # Default paragraph font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    def _add_runs(paragraph, text: str):
        """Parse inline **bold**, *italic*, and `code` markers into runs."""
        token_re = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)')
        for part in token_re.split(text):
            if part.startswith("**") and part.endswith("**") and len(part) > 4:
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*") and len(part) > 2:
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith("`") and part.endswith("`") and len(part) > 2:
                run = paragraph.add_run(part[1:-1])
                run.font.name = "Courier New"
                run.font.size = Pt(10)
            else:
                paragraph.add_run(part)

    for line in markdown_text.splitlines():
        s = line.strip()

        # Headings
        if s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:], level=1)

        # Horizontal rule → paragraph with bottom border
        elif s == "---":
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single")
            bot.set(qn("w:sz"), "6")
            bot.set(qn("w:space"), "1")
            bot.set(qn("w:color"), "999999")
            pBdr.append(bot)
            pPr.append(pBdr)

        # Unordered list
        elif s.startswith("- ") or s.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, s[2:])

        # Ordered list
        elif re.match(r"^\d+\.\s", s):
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, re.sub(r"^\d+\.\s", "", s, count=1))

        # Blank line → paragraph break
        elif s == "":
            doc.add_paragraph()

        # Normal paragraph
        else:
            p = doc.add_paragraph()
            _add_runs(p, s)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# Step 5: Write Output (format-aware, length-aware, dual-path writer)
def step5_write_output(
    research_plan: str,
    initial_research: str,
    web_research: str,
    llm,                               # OllamaLLM — used for model name only
    output_format: str = "blog_post",
    output_length: str = "~1,000 words (2-3 pages)",
    context: str = "",
    resources: str = "",
    num_ctx_writer: int = 32768,
) -> str:
    """Write the final output document in the chosen format and length.

    Primes Ollama with the desired num_ctx via a direct REST call before
    handing off to CrewAI (which uses the bare-string LiteLLM/Ollama path).
    This ensures the model is loaded with the correct context window even
    though CrewAI's LiteLLM layer silently drops extra kwargs like num_ctx.

    Routing:
    - targets < 5,000 words  → single-pass (one Crew call)
    - targets >= 5,000 words → section-by-section (one Crew call per section,
                               results concatenated) to avoid context overflow
    """
    template = FORMAT_TEMPLATES.get(output_format, FORMAT_TEMPLATES["blog_post"])
    format_label = template["label"]
    target_words = _parse_target_word_count(output_length)
    use_sectional = target_words >= 5000

    # Prime Ollama to load the model with the desired context window BEFORE
    # CrewAI starts making calls. This works because Ollama caches the loaded
    # model in memory; subsequent LiteLLM calls from CrewAI will inherit the
    # num_ctx set by this priming request.
    # (CrewAI hardcodes litellm.drop_params=True, so passing num_ctx via
    # kwargs is silently dropped — direct priming is the only reliable path.)
    with st.spinner(f"Preparing writer (context window: {num_ctx_writer:,} tokens) …"):
        _prime_ollama_context(llm.model, num_ctx=num_ctx_writer)

    llm_string = f"ollama/{llm.model}"

    writer_agent = Agent(
        role=template["agent_role"],
        goal=template["agent_goal"],
        backstory=template["agent_backstory"],
        llm=llm_string,        # Bare string — uses the working LiteLLM/Ollama path
        verbose=True,
        allow_delegation=False,
    )

    # ------------------------------------------------------------------
    # PATH A: Section-by-section writing (targets >= 5,000 words)
    # Each section is a separate Crew call with a compressed research
    # summary, keeping each call's input well within the context window.
    # ------------------------------------------------------------------
    if use_sectional:
        st.subheader(f"Step 6: Writing {format_label} (section by section)")
        sections = TEMPLATE_SECTIONS.get(output_format, TEMPLATE_SECTIONS["blog_post"])
        words_per_section = max(400, target_words // len(sections))

        # Compress all research inputs to a fixed-size block (~7 K tokens max).
        # This ensures each section call has a predictable, moderate input size
        # even when the raw research content is very large.
        research_summary = (
            f"RESEARCH PLAN SUMMARY:\n{research_plan[:1500]}\n\n"
            f"INITIAL RESEARCH (KEY POINTS):\n{initial_research[:2000]}\n\n"
            f"WEB RESEARCH FINDINGS:\n{web_research[:2000]}\n"
            f"{context}\n{resources[:1500]}"
        )

        section_outputs: list[str] = []

        for i, section_name in enumerate(sections):
            is_first = i == 0
            is_last  = i == len(sections) - 1

            # Pass the last ~400 words of prior output as a continuity anchor
            # so each section flows naturally from the previous one.
            prior_context = ""
            if section_outputs:
                prior_words = " ".join(section_outputs).split()
                prior_context = (
                    "\n\nPREVIOUS SECTIONS WRITTEN SO FAR "
                    "(last portion shown for continuity — do NOT repeat it):\n"
                    + " ".join(prior_words[-400:])
                    + "\n[Continue seamlessly from here]\n"
                )

            section_task = Task(
                description=f"""
                You are writing ONE section of a {format_label}.

                --- Research Material ---
                {research_summary}
                {prior_context}

                --- Your Task ---
                Write ONLY the section titled: "{section_name}"
                This is section {i + 1} of {len(sections)}.
                {"This is the FIRST section — include the full document title and a compelling opening paragraph." if is_first else ""}
                {"This is the FINAL section — write a strong conclusion that ties the entire document together." if is_last else ""}

                Target length for THIS section: approximately {words_per_section} words.
                Write ONLY this section. Do NOT write any other sections.
                Start with a clear ## heading for "{section_name}".
                Use markdown formatting throughout.
                Cite sources as [Source: URL or Title] where relevant.

                Format guidance: {template["structure_guidance"]}
                """,
                agent=writer_agent,
                expected_output=(
                    f"The '{section_name}' section in markdown "
                    f"(~{words_per_section} words) with a ## heading"
                ),
            )

            section_crew = Crew(
                agents=[writer_agent],
                tasks=[section_task],
                process=Process.sequential,
                verbose=True,
                memory=False,
                cache=False,
                output_log_file=False,
            )

            with st.spinner(
                f"Writing section {i + 1}/{len(sections)}: {section_name} …"
            ):
                section_result = section_crew.kickoff()

            section_outputs.append(str(section_result))

        return "\n\n---\n\n".join(section_outputs)

    # ------------------------------------------------------------------
    # PATH B: Single-pass writing (targets < 5,000 words)
    # ------------------------------------------------------------------
    st.subheader(f"Step 6: Writing {format_label}")
    # Per-section word guide helps the model distribute length evenly
    words_per_section = max(200, target_words // 4)

    writing_task = Task(
        description=f"""
        You are writing a {format_label}.

        --- Research Plan ---
        {research_plan}

        --- Initial Research ---
        {initial_research}

        --- Web Research Findings ---
        {web_research}
        {context}
        {resources}

        --- Format and Structure Instructions ---
        {template["structure_guidance"]}

        --- Length Requirement ---
        Target length: {output_length} (approximately {target_words} words).

        CRITICAL: Write the COMPLETE document in its entirety right now.
        - Do NOT write a partial draft, outline, or introduction only
        - Do NOT stop early or use placeholders like "[section continues...]"
        - Write ALL sections fully from start to finish
        - Aim for at least {words_per_section} words per major section
        - Keep writing until the document is genuinely complete
        If you finish before the word target, expand sections with more
        detail, examples, evidence, and analysis until the target is met.

        --- Citation Instructions ---
        Cite sources using [Source: URL or Title] inline format.
        Prioritise and explicitly cite any custom resources listed above.

        Use markdown formatting throughout. Produce the complete document now.
        """,
        agent=writer_agent,
        expected_output=(
            f"Complete {format_label} in markdown (~{target_words} words) with citations"
        ),
    )

    crew = Crew(
        agents=[writer_agent],
        tasks=[writing_task],
        process=Process.sequential,
        verbose=True,
        memory=False,
        cache=False,
        output_log_file=False,
    )

    with st.spinner(f"Writing {format_label} ({output_length}) …"):
        result = crew.kickoff()

    return str(result)

# ---------------------------------------------------------------------------
# Workflow stage -> display index mapping (for the sidebar indicator)
# Stage values: 0 -> 1 -> 2 -> 3 -> 4 -> 45 -> 5
# ---------------------------------------------------------------------------
_STAGE_ORDER: list[int] = [0, 1, 2, 3, 4, 45, 5]
_STAGE_LABELS: list[str] = [
    "0. Research Context",
    "1. Enter Research Request",
    "2. Research Plan & Initial Research",
    "3. Gap Analysis",
    "4. Web Research",
    "5. Output Format & Length",
    "6. Write & Download Output",
]


def _stage_display_index(stage: int) -> int:
    """Return the 0-based display index for the current workflow stage."""
    try:
        return _STAGE_ORDER.index(stage)
    except ValueError:
        return 0


# ---------------------------------------------------------------------------
# Main Streamlit App
# ---------------------------------------------------------------------------

def main() -> None:
    st.title("🤖 AI Research Agent")
    st.markdown("*Autonomous research agent with human-in-the-loop checkpoints*")

    # ------------------------------------------------------------------
    # Sidebar
    # ------------------------------------------------------------------
    with st.sidebar:
        st.header("⚙️ Configuration")

        # Model selection
        installed_models = get_installed_ollama_models()
        if not installed_models:
            st.error("⚠️ No Ollama models found!")
            st.markdown("Please install a model first:")
            st.code("ollama pull llama2", language="bash")
            st.stop()

        model_name: str = st.selectbox(
            "🤖 Select Ollama Model",
            installed_models,
            help="Select from your locally installed Ollama models",
        )

        if not verify_model_exists(model_name):
            st.error(f"⚠️ Model '{model_name}' not found in Ollama!")
            st.code(f"ollama pull {model_name}", language="bash")
            st.stop()

        with st.expander("ℹ️ Model Information"):
            st.markdown(f"**Selected Model:** `{model_name}`")
            st.markdown("**Status:** ✅ Available")
            try:
                proc = subprocess.run(
                    ["ollama", "list"], capture_output=True, text=True, timeout=5
                )
                if proc.returncode == 0:
                    for line in proc.stdout.split("\n"):
                        if model_name in line:
                            parts = line.split()
                            if len(parts) >= 3:
                                st.markdown(f"**Size:** {parts[2]}")
                            break
            except Exception:
                pass
            st.markdown("---")
            st.markdown("**Your Installed Models:**")
            for m in installed_models:
                prefix = "- ✅ **" if m == model_name else "- "
                suffix = "** (selected)" if m == model_name else ""
                st.markdown(f"{prefix}{m}{suffix}")
            st.markdown("---")
            st.markdown("**Refresh models:** Restart the app after installing new models")
            st.code("ollama pull <model-name>", language="bash")

        # Writer context window setting
        st.markdown("---")
        st.markdown("### 🖥️ Writer Context Window")
        num_ctx_writer: int = st.select_slider(
            "Context window (tokens)",
            options=[8192, 16384, 24576, 32768, 49152, 65536],
            value=st.session_state.get("num_ctx_writer", 32768),
            help=(
                "Sets how many tokens the writer agent can use (input + output combined). "
                "32K handles most long outputs. Reduce to 16K if you have <16 GB VRAM. "
                "Increase to 64K for very long documents if your model supports it. "
                "Changing this will clear any cached written output."
            ),
        )
        if st.session_state.get("num_ctx_writer") != num_ctx_writer:
            # Clear written output so it is regenerated with the new context window
            st.session_state.written_output = None
            st.session_state.blog_post = None

        st.session_state.num_ctx_writer = num_ctx_writer

        # Web search settings
        st.markdown("---")
        st.markdown("### 🌐 Web Search Settings")

        enable_web_search: bool = st.checkbox(
            "Enable Web Search",
            value=True,
            help="Uncheck to skip web research and rely only on local LLM knowledge",
        )
        st.session_state.enable_web_search = enable_web_search

        if enable_web_search:
            search_method: str = st.radio(
                "Search Method",
                options=["selenium", "duckduckgo"],
                format_func=lambda x: (
                    "🌐 Selenium (Headless Browser)"
                    if x == "selenium"
                    else "🦆 DuckDuckGo API"
                ),
                help="Selenium is more reliable but slower. DuckDuckGo is faster but may be rate-limited.",
            )
            st.session_state.search_method = search_method

            if search_method == "selenium":
                browser_type: str = st.selectbox(
                    "Browser",
                    options=["chrome", "firefox"],
                    format_func=lambda x: "🌐 Chrome" if x == "chrome" else "🦊 Firefox",
                )
                st.session_state.browser_type = browser_type

                import platform as _platform
                match _platform.system():
                    case "Darwin":
                        if _platform.machine() == "arm64":
                            st.info(
                                "🍎 **MacOS ARM Detected**\n\n"
                                "**Recommended:** Use Firefox for best compatibility\n\n"
                                "**Chrome issues?** Install ARM version from "
                                "[google.com/chrome](https://www.google.com/chrome/)"
                            )
                        else:
                            st.info("🍎 **MacOS Intel** — Both browsers should work well")
                    case _:
                        st.info(f"💡 Using {browser_type.capitalize()} in headless mode")
            else:
                st.info("💡 Using DuckDuckGo API (may be rate-limited)")
        else:
            st.warning("⚠️ Web search disabled. Research will rely only on LLM knowledge.")

        # Persistent resource panel
        render_resource_panel()

        # Workflow stage indicator
        st.markdown("---")
        st.markdown("### 📋 Workflow Stages")
        current_display = _stage_display_index(st.session_state.workflow_stage)
        for i, label in enumerate(_STAGE_LABELS):
            if i < current_display:
                st.markdown(f"~~{label}~~ ✓")
            elif i == current_display:
                st.markdown(f"**→ {label}**")
            else:
                st.markdown(label)

        st.markdown("---")
        if st.button("🔄 Reset Workflow"):
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.rerun()

    # ------------------------------------------------------------------
    # Initialise LLM (cached)
    # ------------------------------------------------------------------
    llm = get_llm(model_name)

    # ------------------------------------------------------------------
    # Stage 0 -- Research Context + Topic Entry
    # ------------------------------------------------------------------
    if st.session_state.workflow_stage == 0:
        st.header("Step 0: Research Context")
        st.markdown(
            "Configure who this research is for and what perspective it should take. "
            "These settings shape every stage — from the research plan to the final document."
        )

        PERSPECTIVE_OPTIONS = [
            "Neutral / Objective",
            "Skeptical Scientist",
            "Science Communicator",
            "Policy Advocate",
            "Industry Practitioner",
            "Educator / Teacher",
            "Investor / Business Strategist",
            "Custom (specify below)",
        ]
        AUDIENCE_OPTIONS = [
            "General Public",
            "Graduate / Doctoral Students",
            "Domain Experts / Researchers",
            "Policymakers / Government Officials",
            "Funding Agencies (NSF / NIH / DOE)",
            "Industry Professionals",
            "Undergraduate Students",
            "Custom (specify below)",
        ]

        col_p, col_a = st.columns(2)
        with col_p:
            perspective_choice: str = st.selectbox(
                "Perspective / Voice",
                PERSPECTIVE_OPTIONS,
                help="The viewpoint from which the research is framed.",
            )
            if perspective_choice == "Custom (specify below)":
                perspective_value: str = st.text_input(
                    "Custom perspective", placeholder="e.g. Climate policy advocate"
                )
            else:
                perspective_value = perspective_choice

        with col_a:
            audience_choice: str = st.selectbox(
                "Target Audience",
                AUDIENCE_OPTIONS,
                help="Who will read the final output.",
            )
            if audience_choice == "Custom (specify below)":
                audience_value: str = st.text_input(
                    "Custom audience", placeholder="e.g. High-school science teachers"
                )
            else:
                audience_value = audience_choice

        framing_notes: str = st.text_area(
            "Framing Notes (optional)",
            placeholder=(
                "Any additional framing guidance, e.g. 'Assume reader is sceptical of AI claims' "
                "or 'Emphasise economic impacts over technical details'."
            ),
            height=80,
        )

        st.divider()

        if st.session_state.context_configured:
            st.info(
                f"**Current context:** Perspective: *{st.session_state.perspective}* | "
                f"Audience: *{st.session_state.audience}*"
            )

        st.subheader("Research Topic")
        user_input: str = st.text_area(
            "What would you like to research?",
            value=st.session_state.user_input,
            height=120,
            placeholder=(
                "Example: Research the latest developments in quantum computing "
                "and their potential applications in cryptography"
            ),
        )

        if st.button("🚀 Start Research", type="primary"):
            if not user_input.strip():
                st.warning("Please enter a research topic.")
            else:
                st.session_state.perspective = perspective_value
                st.session_state.audience = audience_value
                st.session_state.framing_notes = framing_notes
                st.session_state.context_configured = True
                st.session_state.user_input = user_input
                st.session_state.workflow_stage = 1
                st.rerun()

    # ------------------------------------------------------------------
    # Stage 1 -- Research Plan + Initial Research
    # ------------------------------------------------------------------
    elif st.session_state.workflow_stage == 1:
        st.header("Step 1 & 2: Research Plan and Initial Research")

        if st.session_state.perspective or st.session_state.audience:
            st.info(
                f"**Perspective:** {st.session_state.perspective or '—'}  |  "
                f"**Audience:** {st.session_state.audience or '—'}"
            )

        context = build_context_block()
        resources = build_resources_block()

        if st.session_state.research_plan is None:
            st.session_state.research_plan = step1_interpret_and_plan(
                st.session_state.user_input, llm, context=context, resources=resources
            )

        st.subheader("Research Plan")
        st.text_area("Plan", st.session_state.research_plan, height=300)

        if st.session_state.initial_research is None:
            st.session_state.initial_research = step2_initial_research(
                st.session_state.research_plan, llm, context=context, resources=resources
            )

        st.subheader("Initial Research")
        st.text_area("Research", st.session_state.initial_research, height=400)

        col1, col2, col3 = st.columns(3)
        with col1:
            if st.button("✅ Proceed to Gap Analysis", type="primary"):
                st.session_state.workflow_stage = 2
                st.rerun()
        with col2:
            if st.button("🔄 Redo Research Plan"):
                st.session_state.research_plan = None
                st.session_state.initial_research = None
                st.rerun()
        with col3:
            if st.button("✏️ Edit Context"):
                st.session_state.workflow_stage = 0
                st.rerun()

    # ------------------------------------------------------------------
    # Stage 2 -- Gap Analysis
    # ------------------------------------------------------------------
    elif st.session_state.workflow_stage == 2:
        st.header("Step 3: Gap Analysis")

        if st.session_state.perspective or st.session_state.audience:
            st.info(
                f"**Perspective:** {st.session_state.perspective or '—'}  |  "
                f"**Audience:** {st.session_state.audience or '—'}"
            )

        context = build_context_block()
        resources = build_resources_block()

        if st.session_state.gap_analysis is None:
            st.session_state.gap_analysis = step3_gap_analysis(
                st.session_state.research_plan,
                st.session_state.initial_research,
                llm,
                context=context,
                resources=resources,
            )

        st.text_area("Gap Analysis", st.session_state.gap_analysis, height=400)

        col1, col2 = st.columns(2)
        with col1:
            if st.button("✅ Proceed to Web Research", type="primary"):
                st.session_state.workflow_stage = 3
                st.rerun()
        with col2:
            if st.button("🔄 Redo Gap Analysis"):
                st.session_state.gap_analysis = None
                st.rerun()

    # ------------------------------------------------------------------
    # Stage 3 -- Web Research
    # ------------------------------------------------------------------
    elif st.session_state.workflow_stage == 3:
        st.header("Step 4: Web Research")

        if st.session_state.perspective or st.session_state.audience:
            st.info(
                f"**Perspective:** {st.session_state.perspective or '—'}  |  "
                f"**Audience:** {st.session_state.audience or '—'}"
            )

        context = build_context_block()
        resources = build_resources_block()

        if not st.session_state.get("enable_web_search", True):
            st.warning("⚠️ Web search is disabled. Skipping to output format selection.")
            st.session_state.web_research = (
                "Web search was disabled. Proceeding with information from "
                "initial research and gap analysis only."
            )
            col1, col2 = st.columns(2)
            with col1:
                if st.button("✅ Proceed to Output Format", type="primary"):
                    st.session_state.workflow_stage = 45
                    st.rerun()
            with col2:
                if st.button("🔙 Back to Gap Analysis"):
                    st.session_state.workflow_stage = 2
                    st.rerun()
        else:
            if st.session_state.web_research is None:
                st.session_state.web_research = step4_web_research(
                    st.session_state.gap_analysis,
                    llm,
                    context=context,
                    resources=resources,
                )

            st.text_area("Web Research Findings", st.session_state.web_research, height=400)

            col1, col2 = st.columns(2)
            with col1:
                if st.button("✅ Proceed to Output Format", type="primary"):
                    st.session_state.workflow_stage = 45
                    st.rerun()
            with col2:
                if st.button("🔄 Redo Web Research"):
                    st.session_state.web_research = None
                    st.rerun()

    # ------------------------------------------------------------------
    # Stage 45 -- Output Format & Length Selection
    # ------------------------------------------------------------------
    elif st.session_state.workflow_stage == 45:
        st.header("Step 5: Choose Output Format & Length")
        st.markdown(
            "All research is complete. Select how you want the final document structured "
            "and how long it should be."
        )

        if st.session_state.perspective or st.session_state.audience:
            st.info(
                f"**Perspective:** {st.session_state.perspective or '—'}  |  "
                f"**Audience:** {st.session_state.audience or '—'}"
            )

        FORMAT_DISPLAY: dict[str, str] = {k: v["label"] for k, v in FORMAT_TEMPLATES.items()}
        LENGTH_OPTIONS: list[str] = [
            "~500 words (1 page)",
            "~1,000 words (2-3 pages)",
            "~2,500 words (5 pages)",
            "~5,000 words (10 pages)",
            "Custom…",
        ]

        col_fmt, col_len = st.columns(2)
        with col_fmt:
            format_key: str = st.selectbox(
                "Output Format",
                list(FORMAT_DISPLAY.keys()),
                format_func=lambda k: FORMAT_DISPLAY[k],
                index=list(FORMAT_DISPLAY.keys()).index(
                    st.session_state.output_format
                    if st.session_state.output_format in FORMAT_DISPLAY
                    else "blog_post"
                ),
            )
            st.caption(FORMAT_TEMPLATES[format_key]["structure_guidance"])

        with col_len:
            length_choice: str = st.selectbox(
                "Target Length",
                LENGTH_OPTIONS,
                index=LENGTH_OPTIONS.index(st.session_state.output_length)
                if st.session_state.output_length in LENGTH_OPTIONS
                else 1,
            )
            if length_choice == "Custom…":
                custom_length: str = st.text_input(
                    "Custom length target",
                    placeholder="e.g. 3,500 words or 7 pages",
                )
                final_length = custom_length if custom_length.strip() else "~1,000 words"
            else:
                final_length = length_choice

        st.markdown("---")
        if st.button("✅ Generate Output", type="primary"):
            st.session_state.output_format = format_key
            st.session_state.output_length = final_length
            st.session_state.output_format_selected = True
            st.session_state.workflow_stage = 5
            st.rerun()

        if st.button("🔙 Back to Web Research"):
            st.session_state.workflow_stage = 3
            st.rerun()

    # ------------------------------------------------------------------
    # Stage 5 -- Write Output
    # ------------------------------------------------------------------
    elif st.session_state.workflow_stage == 5:
        fmt_label = FORMAT_TEMPLATES.get(
            st.session_state.output_format, FORMAT_TEMPLATES["blog_post"]
        )["label"]
        st.header(f"Step 6: Writing — {fmt_label}")

        if st.session_state.perspective or st.session_state.audience:
            st.info(
                f"**Perspective:** {st.session_state.perspective or '—'}  |  "
                f"**Audience:** {st.session_state.audience or '—'}  |  "
                f"**Length:** {st.session_state.output_length}"
            )

        context = build_context_block()
        resources = build_resources_block()

        if st.session_state.written_output is None:
            st.session_state.written_output = step5_write_output(
                research_plan=st.session_state.research_plan,
                initial_research=st.session_state.initial_research,
                web_research=st.session_state.web_research,
                llm=llm,
                output_format=st.session_state.output_format,
                output_length=st.session_state.output_length,
                context=context,
                resources=resources,
                num_ctx_writer=st.session_state.get("num_ctx_writer", 32768),
            )
            # Keep legacy alias so HTML step can access it
            st.session_state.blog_post = st.session_state.written_output

        st.markdown(f"### {fmt_label} Preview")
        st.markdown(st.session_state.written_output)

        safe_fmt = st.session_state.output_format.replace(" ", "_")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        # Generate DOCX bytes deterministically (no LLM needed)
        docx_bytes = _markdown_to_docx(st.session_state.written_output, title=fmt_label)

        col1, col2, col3 = st.columns(3)
        with col1:
            st.download_button(
                "📥 Download DOCX",
                docx_bytes,
                file_name=f"{safe_fmt}_{timestamp}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
            )
        with col2:
            st.download_button(
                "📥 Download Markdown",
                st.session_state.written_output,
                file_name=f"{safe_fmt}_{timestamp}.md",
                mime="text/markdown",
            )
        with col3:
            if st.button("🔄 Redo Output"):
                st.session_state.written_output = None
                st.session_state.blog_post = None
                st.rerun()

        st.markdown("---")
        st.success("✅ Workflow Complete! Download your document above.")

        col_back, col_new = st.columns(2)
        with col_back:
            if st.button("🔙 Change Format / Length"):
                st.session_state.written_output = None
                st.session_state.blog_post = None
                st.session_state.workflow_stage = 45
                st.rerun()
        with col_new:
            if st.button("🎉 Start New Research"):
                for key in list(st.session_state.keys()):
                    del st.session_state[key]
                st.rerun()



if __name__ == "__main__":
    main()
